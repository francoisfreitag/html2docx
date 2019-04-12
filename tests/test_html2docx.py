import json
import pathlib

import docx
import pytest

from html2docx import html2docx


def generate_testdata():
    datadir = pathlib.Path(__file__).parent.joinpath("data")
    for html_path in datadir.glob("*.html"):
        spec_path = datadir.joinpath(f"{html_path.stem}.json")
        yield html_path, spec_path


@pytest.mark.parametrize("html_path,spec_path", generate_testdata())
def test_html2docx(html_path, spec_path):
    title = html_path.name
    with html_path.open() as fp:
        html = fp.read()
    with spec_path.open() as fp:
        spec = json.load(fp)

    buf = html2docx(html, title=title)
    doc = docx.Document(buf)

    assert doc.core_properties.title == title
    assert len(doc.paragraphs) == len(spec)
    for p, p_spec in zip(doc.paragraphs, spec):
        assert p.text == p_spec["text"]
        assert p.style.name == p_spec.get("style", "Normal")

        runs_spec = p_spec["runs"]
        assert len(p.runs) == len(runs_spec)
        for run, run_spec in zip(p.runs, runs_spec):
            assert run.text == run_spec["text"]
            assert run.bold is run_spec.get(
                "bold", False
            ), f"Wrong bold for text '{run.text}'."
            assert run.italic is run_spec.get(
                "italic", False
            ), f"Wrong italic for text '{run.text}'."
            assert run.underline is run_spec.get(
                "underline", False
            ), f"Wrong underline for text '{run.text}'."


def test_href():
    html = '<a href="https://github.com/">GitHub</a>'
    buf = html2docx(html, title="test")
    doc = docx.Document(buf)

    assert len(doc.paragraphs) == 1

    p = doc.paragraphs[0]
    children = p._p.getchildren()
    assert len(children) == 1

    child = children[0]
    assert (
        child.tag
        == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hyperlink"
    )

    runs = child.getchildren()
    assert len(runs) == 1

    run = runs[0]
    run_components = run.getchildren()
    assert len(run_components) == 2

    rPr, text_elt = run_components
    assert (
        rPr.tag == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPr"
    )
    assert text_elt.text == "GitHub"

    style_elts = rPr.getchildren()
    assert len(style_elts) == 2

    underline, color = style_elts
    assert (
        underline.tag
        == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}u"
    )
    assert underline.values() == ["single"]
    assert (
        color.tag
        == "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}color"
    )
    assert color.values() == ["0000FF"]
